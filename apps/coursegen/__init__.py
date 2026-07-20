import json
import io
import re
import html as html_lib
from flask import Blueprint, render_template, request, jsonify, Response, stream_with_context
from models import db, Setting, Course, CourseModule, CourseQuiz, CourseAssignment

coursegen_bp = Blueprint("coursegen", __name__, url_prefix="/coursegen")

import shutil as _shutil
WKHTMLTOPDF = (
    _shutil.which("wkhtmltopdf")
    or "/usr/bin/wkhtmltopdf"           # PythonAnywhere default
    or "/nix/store/hxiay4lkq4389vxnhnb3d0pbaw6siwkw-wkhtmltopdf/bin/wkhtmltopdf"
)


# ── Helpers ───────────────────────────────────────────────────────────────────

def get_groq_key():
    s = Setting.query.filter_by(key="GROQ_API_KEY").first()
    return s.value.strip() if s else ""


def extract_json(text: str) -> dict:
    """Robustly extract the first complete JSON object from arbitrary text."""
    text = text.strip()

    # 1. Direct parse
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass

    # 2. Strip markdown code fence
    m = re.search(r'```(?:json)?\s*([\s\S]*?)```', text)
    if m:
        try:
            return json.loads(m.group(1).strip())
        except json.JSONDecodeError:
            pass

    # 3. Find outermost { … } by tracking brace depth
    start = text.find('{')
    if start != -1:
        depth = 0
        in_str = False
        escape = False
        for i in range(start, len(text)):
            ch = text[i]
            if escape:
                escape = False
                continue
            if ch == '\\' and in_str:
                escape = True
                continue
            if ch == '"':
                in_str = not in_str
                continue
            if in_str:
                continue
            if ch == '{':
                depth += 1
            elif ch == '}':
                depth -= 1
                if depth == 0:
                    try:
                        return json.loads(text[start:i + 1])
                    except json.JSONDecodeError:
                        break

    raise json.JSONDecodeError(
        f"No valid JSON object found. Raw response (first 300 chars): {text[:300]}",
        text, 0
    )


def groq_json(client, messages, max_tokens=2000, retries=5):
    """Call Groq and robustly parse the JSON response.
    Retries on 429 rate-limit errors (exponential back-off) and on JSON
    decode failures (up to 2 extra attempts with a nudge message)."""
    import time
    from groq import RateLimitError

    delay = 1.5
    json_failures = 0

    for attempt in range(retries):
        try:
            resp = client.chat.completions.create(
                model="llama-3.1-8b-instant",
                messages=messages,
                max_tokens=max_tokens,
                temperature=0.6,
                response_format={"type": "json_object"},
            )
            raw = resp.choices[0].message.content or ""
            return extract_json(raw)
        except RateLimitError as e:
            if attempt == retries - 1:
                raise
            wait = delay
            m = re.search(r'in (\d+(?:\.\d+)?)s', str(e))
            if m:
                wait = float(m.group(1)) + 0.5
            else:
                m = re.search(r'in (\d+)ms', str(e))
                if m:
                    wait = int(m.group(1)) / 1000 + 0.5
            time.sleep(wait)
            delay *= 2
        except json.JSONDecodeError:
            json_failures += 1
            if json_failures >= 3 or attempt == retries - 1:
                raise
            # Nudge the model to return clean JSON on the next attempt
            messages = list(messages) + [
                {"role": "assistant", "content": raw},
                {"role": "user",      "content": "Your response was not valid JSON. Return only a valid JSON object, nothing else."},
            ]
            time.sleep(1.0)


def sse(data: dict) -> str:
    return f"data: {json.dumps(data)}\n\n"


# ── Markdown → HTML (for PDF) ─────────────────────────────────────────────────

def _inline(text):
    text = html_lib.escape(text)
    text = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', text)
    text = re.sub(r'\*(.+?)\*',     r'<em>\1</em>',         text)
    text = re.sub(r'`([^`]+)`',     r'<code>\1</code>',     text)
    return text


def md_to_html(text):
    lines   = text.split('\n')
    out     = []
    in_ul   = False
    in_pre  = False
    pre_buf = []

    def flush_ul():
        nonlocal in_ul
        if in_ul:
            out.append('</ul>')
            in_ul = False

    for line in lines:
        if line.startswith('```'):
            if in_pre:
                out.append('<pre><code>' + html_lib.escape('\n'.join(pre_buf)) + '</code></pre>')
                pre_buf.clear()
                in_pre = False
            else:
                flush_ul()
                in_pre = True
            continue
        if in_pre:
            pre_buf.append(line)
            continue
        if line.startswith('#### '):
            flush_ul(); out.append(f'<h4>{_inline(line[5:])}</h4>')
        elif line.startswith('### '):
            flush_ul(); out.append(f'<h3>{_inline(line[4:])}</h3>')
        elif line.startswith('## '):
            flush_ul(); out.append(f'<h2>{_inline(line[3:])}</h2>')
        elif line.startswith('# '):
            flush_ul(); out.append(f'<h2>{_inline(line[2:])}</h2>')
        elif re.match(r'^[-*+] ', line):
            if not in_ul:
                out.append('<ul>'); in_ul = True
            out.append(f'<li>{_inline(line[2:])}</li>')
        elif line.strip() == '':
            flush_ul(); out.append('')
        else:
            flush_ul(); out.append(f'<p>{_inline(line)}</p>')

    flush_ul()
    return '\n'.join(out)


# ── PDF builder ───────────────────────────────────────────────────────────────

PDF_CSS = """
body{font-family:Arial,sans-serif;color:#1a1a1a;padding:36px 48px;max-width:820px;margin:0 auto}
h1{color:#4f46e5;font-size:24px;margin-bottom:4px}
h2{color:#312e81;font-size:18px;border-bottom:2px solid #e0e7ff;padding-bottom:6px;margin-top:28px}
h3{color:#3730a3;font-size:15px;margin-top:18px}
h4{color:#4338ca;font-size:14px;margin-top:14px}
p{line-height:1.7;margin:8px 0;font-size:14px}
ul{padding-left:22px;margin:8px 0}
li{margin:4px 0;font-size:14px;line-height:1.6}
code{background:#f3f4f6;padding:1px 5px;border-radius:3px;font-family:monospace;font-size:12px;color:#6366f1}
pre{background:#f3f4f6;padding:14px;border-radius:6px;overflow-x:auto;font-size:12px;margin:12px 0}
pre code{background:none;padding:0;color:#374151}
.meta{color:#6b7280;font-size:13px;margin-bottom:24px}
.quiz-block{margin:10px 0 16px;padding:14px;background:#f9fafb;border-left:3px solid #6366f1;border-radius:4px}
.quiz-q{font-weight:700;font-size:14px;margin-bottom:8px;color:#1f2937}
.quiz-opt{font-size:13px;margin:4px 0;padding:4px 0;color:#374151}
.quiz-opt.correct{color:#059669;font-weight:700}
.asgn{background:#fffbeb;border-left:3px solid #f59e0b;padding:14px;border-radius:4px;margin-top:10px}
.asgn-title{font-weight:800;font-size:15px;color:#92400e;margin-bottom:6px}
.page-break{page-break-after:always}
.module-header{background:#4f46e5;color:#fff;padding:12px 16px;border-radius:6px;margin-bottom:20px}
.module-header h2{color:#fff;border:none;margin:0;font-size:17px}
.module-header .num{font-size:11px;opacity:.8;margin-bottom:4px;text-transform:uppercase;letter-spacing:.05em}
"""

def build_module_html(course, mod, show_answers=True):
    letters = ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H']
    quiz_html = ""
    if mod.quizzes:
        quiz_html = "<h2>Quiz</h2>"
        for i, q in enumerate(mod.quizzes):
            opts = json.loads(q.options)
            opts_html = ""
            for j, opt in enumerate(opts):
                letter = letters[j] if j < len(letters) else str(j)
                is_correct = (j == q.correct_answer) and show_answers
                cls = ' class="quiz-opt correct"' if is_correct else ' class="quiz-opt"'
                mark = " ✓" if is_correct else ""
                opts_html += f'<div{cls}>{letter}. {html_lib.escape(opt)}{mark}</div>'
            quiz_html += f'''<div class="quiz-block">
  <div class="quiz-q">{i+1}. {html_lib.escape(q.question)}</div>
  {opts_html}
</div>'''

    asgn_html = ""
    if mod.assignments:
        a = mod.assignments[0]
        asgn_html = f'''<div class="asgn">
  <div class="asgn-title">📝 {html_lib.escape(a.title)}</div>
  <p style="font-size:13px;color:#78350f;line-height:1.7">{html_lib.escape(a.description)}</p>
</div>'''

    content_html = md_to_html(mod.content)

    return f"""<!DOCTYPE html><html><head><meta charset="UTF-8">
<style>{PDF_CSS}</style></head><body>
<h1>{html_lib.escape(course.title)}</h1>
<p class="meta">{html_lib.escape(course.difficulty)} &bull; {html_lib.escape(course.language)} &bull; Module {mod.number} of {course.total_modules}</p>
<div class="module-header">
  <div class="num">Module {mod.number}</div>
  <h2>{html_lib.escape(mod.title)}</h2>
</div>
{content_html}
{quiz_html}
{asgn_html}
</body></html>"""


def build_course_html(course):
    letters = ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H']
    modules_html = ""
    for idx, mod in enumerate(course.modules):
        quiz_html = ""
        if mod.quizzes:
            quiz_html = "<h2>Quiz</h2>"
            for i, q in enumerate(mod.quizzes):
                opts = json.loads(q.options)
                opts_html = "".join(
                    f'<div class="quiz-opt{" correct" if j == q.correct_answer else ""}">'
                    f'{letters[j] if j < len(letters) else j}. {html_lib.escape(opt)}'
                    f'{"  ✓" if j == q.correct_answer else ""}</div>'
                    for j, opt in enumerate(opts)
                )
                quiz_html += f'<div class="quiz-block"><div class="quiz-q">{i+1}. {html_lib.escape(q.question)}</div>{opts_html}</div>'

        asgn_html = ""
        if mod.assignments:
            a = mod.assignments[0]
            asgn_html = f'<div class="asgn"><div class="asgn-title">📝 {html_lib.escape(a.title)}</div><p style="font-size:13px;color:#78350f;line-height:1.7">{html_lib.escape(a.description)}</p></div>'

        pb = '<div class="page-break"></div>' if idx < len(course.modules) - 1 else ''
        modules_html += f"""
<div class="module-header">
  <div class="num">Module {mod.number}</div>
  <h2>{html_lib.escape(mod.title)}</h2>
</div>
{md_to_html(mod.content)}
{quiz_html}
{asgn_html}
{pb}"""

    return f"""<!DOCTYPE html><html><head><meta charset="UTF-8">
<style>{PDF_CSS}</style></head><body>
<h1>{html_lib.escape(course.title)}</h1>
<p class="meta">{html_lib.escape(course.difficulty)} &bull; {html_lib.escape(course.language)} &bull; {course.total_modules} Modules</p>
{modules_html}
</body></html>"""


def html_to_pdf(html_str: str) -> bytes:
    import pdfkit
    config = pdfkit.configuration(wkhtmltopdf=WKHTMLTOPDF)
    options = {
        'quiet': '',
        'encoding': 'UTF-8',
        'margin-top':    '15mm',
        'margin-bottom': '15mm',
        'margin-left':   '15mm',
        'margin-right':  '15mm',
        'enable-local-file-access': '',
    }
    return pdfkit.from_string(html_str, False, options=options, configuration=config)


# ── DOCX builder ──────────────────────────────────────────────────────────────

def build_module_docx(course, mod) -> io.BytesIO:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml.ns import qn
    import docx.oxml as oxml

    doc = Document()

    # Styles
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)

    def set_heading_color(para, r, g, b):
        for run in para.runs:
            run.font.color.rgb = RGBColor(r, g, b)

    # Course title
    t = doc.add_heading(course.title, level=0)
    set_heading_color(t, 79, 70, 229)

    # Meta
    meta = doc.add_paragraph()
    meta.add_run(f"{course.difficulty} · {course.language} · Module {mod.number} of {course.total_modules}").font.color.rgb = RGBColor(107, 114, 128)

    # Module heading
    mh = doc.add_heading(f"Module {mod.number}: {mod.title}", level=1)
    set_heading_color(mh, 49, 46, 129)

    doc.add_paragraph()

    # Content — parse markdown
    for line in mod.content.split('\n'):
        if line.startswith('```'):
            continue
        elif line.startswith('#### '):
            h = doc.add_heading(line[5:], level=4)
            set_heading_color(h, 67, 56, 202)
        elif line.startswith('### '):
            h = doc.add_heading(line[4:], level=3)
            set_heading_color(h, 55, 48, 163)
        elif line.startswith('## '):
            h = doc.add_heading(line[3:], level=2)
            set_heading_color(h, 49, 46, 129)
        elif line.startswith('# '):
            h = doc.add_heading(line[2:], level=2)
            set_heading_color(h, 49, 46, 129)
        elif re.match(r'^[-*+] ', line):
            clean = re.sub(r'\*\*(.+?)\*\*', r'\1', line[2:])
            clean = re.sub(r'\*(.+?)\*',     r'\1', clean)
            clean = re.sub(r'`(.+?)`',        r'\1', clean)
            doc.add_paragraph(clean, style='List Bullet')
        elif line.strip():
            clean = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
            clean = re.sub(r'\*(.+?)\*',     r'\1', clean)
            clean = re.sub(r'`(.+?)`',        r'\1', clean)
            doc.add_paragraph(clean)

    # Quiz
    if mod.quizzes:
        doc.add_paragraph()
        qh = doc.add_heading('Quiz', level=2)
        set_heading_color(qh, 49, 46, 129)
        letters = ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H']
        for i, q in enumerate(mod.quizzes):
            qp = doc.add_paragraph()
            r = qp.add_run(f"{i+1}. {q.question}")
            r.bold = True
            opts = json.loads(q.options)
            for j, opt in enumerate(opts):
                letter = letters[j] if j < len(letters) else str(j)
                is_correct = (j == q.correct_answer)
                op = doc.add_paragraph(style='List Bullet')
                run = op.add_run(f"{letter}. {opt}{' ✓' if is_correct else ''}")
                if is_correct:
                    run.font.color.rgb = RGBColor(5, 150, 105)
                    run.bold = True

    # Assignment
    if mod.assignments:
        doc.add_paragraph()
        ah = doc.add_heading('Assignment', level=2)
        set_heading_color(ah, 49, 46, 129)
        a = mod.assignments[0]
        at = doc.add_heading(a.title, level=3)
        set_heading_color(at, 146, 64, 14)
        doc.add_paragraph(a.description)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def build_course_docx(course) -> io.BytesIO:
    from docx import Document
    from docx.shared import Pt, RGBColor
    letters = ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H']

    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    def set_heading_color(para, r, g, b):
        for run in para.runs:
            run.font.color.rgb = RGBColor(r, g, b)

    t = doc.add_heading(course.title, level=0)
    set_heading_color(t, 79, 70, 229)

    meta = doc.add_paragraph()
    meta.add_run(f"{course.difficulty} · {course.language} · {course.total_modules} Modules").font.color.rgb = RGBColor(107, 114, 128)

    for idx, mod in enumerate(course.modules):
        if idx > 0:
            doc.add_page_break()

        mh = doc.add_heading(f"Module {mod.number}: {mod.title}", level=1)
        set_heading_color(mh, 49, 46, 129)
        doc.add_paragraph()

        for line in mod.content.split('\n'):
            if line.startswith('```'):
                continue
            elif line.startswith('#### '):
                h = doc.add_heading(line[5:], level=4); set_heading_color(h, 67, 56, 202)
            elif line.startswith('### '):
                h = doc.add_heading(line[4:], level=3); set_heading_color(h, 55, 48, 163)
            elif line.startswith('## '):
                h = doc.add_heading(line[3:], level=2); set_heading_color(h, 49, 46, 129)
            elif line.startswith('# '):
                h = doc.add_heading(line[2:], level=2); set_heading_color(h, 49, 46, 129)
            elif re.match(r'^[-*+] ', line):
                clean = re.sub(r'[*`]', '', line[2:])
                doc.add_paragraph(clean, style='List Bullet')
            elif line.strip():
                clean = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
                clean = re.sub(r'[*`]', '', clean)
                doc.add_paragraph(clean)

        if mod.quizzes:
            doc.add_paragraph()
            qh = doc.add_heading('Quiz', level=2); set_heading_color(qh, 49, 46, 129)
            for i, q in enumerate(mod.quizzes):
                qp = doc.add_paragraph()
                qp.add_run(f"{i+1}. {q.question}").bold = True
                opts = json.loads(q.options)
                for j, opt in enumerate(opts):
                    letter = letters[j] if j < len(letters) else str(j)
                    is_correct = (j == q.correct_answer)
                    op = doc.add_paragraph(style='List Bullet')
                    run = op.add_run(f"{letter}. {opt}{' ✓' if is_correct else ''}")
                    if is_correct:
                        run.font.color.rgb = RGBColor(5, 150, 105)
                        run.bold = True

        if mod.assignments:
            doc.add_paragraph()
            ah = doc.add_heading('Assignment', level=2); set_heading_color(ah, 49, 46, 129)
            a = mod.assignments[0]
            at = doc.add_heading(a.title, level=3); set_heading_color(at, 146, 64, 14)
            doc.add_paragraph(a.description)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── Routes ─────────────────────────────────────────────────────────────────────

@coursegen_bp.route("/")
def index():
    courses = Course.query.order_by(Course.created_at.desc()).limit(20).all()
    return render_template("coursegen/index.html", courses=courses)


@coursegen_bp.route("/generate", methods=["POST"])
def generate():
    data             = request.get_json()
    topic            = (data.get("topic") or "").strip()
    num_modules      = max(2, min(50, int(data.get("num_modules", 3))))
    difficulty       = data.get("difficulty", "Beginner")
    language         = (data.get("language") or "English").strip() or "English"
    words_per_module = max(100, min(1500, int(data.get("words_per_module", 300))))
    quiz_count       = max(1, min(15, int(data.get("quiz_count", 3))))
    include_assign   = data.get("include_assignment", True)

    if not topic:
        return jsonify({"error": "Topic is required."}), 400

    key = get_groq_key()
    if not key:
        return jsonify({"error": "Groq API key not configured. Set it in Admin → Settings."}), 400

    def stream():
        from groq import Groq
        from app import app as flask_app

        with flask_app.app_context():
            try:
                client = Groq(api_key=key)

                # ── Phase 1: outline ─────────────────────────────────────────
                yield sse({"type": "status", "msg": "Generating course outline…"})

                outline = groq_json(client, [
                    {"role": "system", "content": (
                        "You are an expert course designer. "
                        "Always respond with valid JSON only — no markdown, no explanation, no extra text."
                    )},
                    {"role": "user", "content": (
                        f"Create a course outline for a {difficulty}-level course on the topic: \"{topic}\". "
                        f"The course must have exactly {num_modules} module(s). Language: {language}.\n\n"
                        "Respond with a single JSON object in this exact shape:\n"
                        "{\n"
                        '  "title": "<course title>",\n'
                        '  "description": "<2-sentence overview>",\n'
                        '  "modules": [\n'
                        '    {"number": 1, "title": "<module title>"},\n'
                        f'    ... (exactly {num_modules} items, numbered 1 to {num_modules})\n'
                        "  ]\n"
                        "}\n\n"
                        "Rules: output ONLY the JSON object, nothing else."
                    )},
                ], max_tokens=150 + num_modules * 30)

                course = Course(
                    title=outline.get("title", topic),
                    topic=topic,
                    difficulty=difficulty,
                    language=language,
                    total_modules=num_modules,
                )
                db.session.add(course)
                db.session.flush()

                module_list = outline.get("modules", [])
                # Ensure we have exactly num_modules entries
                while len(module_list) < num_modules:
                    n = len(module_list) + 1
                    module_list.append({"number": n, "title": f"Module {n}"})

                yield sse({
                    "type":    "outline",
                    "title":   course.title,
                    "modules": module_list,
                })

                # ── Phase 2: each module ─────────────────────────────────────
                assign_instruction = (
                    ',"assignment":{"title":"Assignment title","description":"Detailed assignment instructions, at least 80 words."}'
                    if include_assign else ""
                )
                # tokens ≈ words × 1.4; JSON overhead + quiz + assignment add ~600
                max_out = max(2000, int(words_per_module * 1.5 * 4) + quiz_count * 200 + 800)

                for i, mod_info in enumerate(module_list[:num_modules]):
                    mod_num   = mod_info.get("number", i + 1)
                    mod_title = mod_info.get("title", f"Module {mod_num}")

                    yield sse({
                        "type":    "progress",
                        "current": i + 1,
                        "total":   num_modules,
                        "msg":     f"Generating module {mod_num}: {mod_title}…",
                    })

                    quiz_shape = (
                        '    {"question": "<question text>?", "options": ["<A>", "<B>", "<C>", "<D>"], "correct": <0-based index>}'
                    )
                    assign_shape = (
                        ',\n  "assignment": {"title": "<assignment title>", "description": "<instructions, min 80 words>"}'
                        if include_assign else ""
                    )
                    mod_data = groq_json(client, [
                        {"role": "system", "content": (
                            "You are an expert course designer. "
                            "Always respond with valid JSON only — no markdown, no explanation, no extra text."
                        )},
                        {"role": "user", "content": (
                            f"Generate the full content for Module {mod_num}: \"{mod_title}\".\n"
                            f"Course topic: \"{topic}\" | Difficulty: {difficulty} | Language: {language}\n\n"
                            f"Requirements:\n"
                            f"- content: at least {words_per_module} words of rich educational material "
                            f"using ## headings, bullet points, numbered lists, and code blocks where relevant.\n"
                            f"- quiz: exactly {quiz_count} multiple-choice question(s), each with exactly 4 options "
                            f"and a 0-based integer 'correct' index.\n"
                            + ("- assignment: one practical assignment with a title and detailed description (min 80 words).\n" if include_assign else "")
                            + "\nRespond with a single JSON object in this exact shape:\n"
                            "{\n"
                            '  "content": "<full markdown content>",\n'
                            f'  "quiz": [\n{quiz_shape},\n    ... ({quiz_count} items total)\n  ]'
                            + assign_shape
                            + "\n}\n\n"
                            "Rules: output ONLY the JSON object, nothing else."
                        )},
                    ], max_tokens=max_out)

                    mod = CourseModule(
                        course_id=course.id,
                        number=mod_num,
                        title=mod_title,
                        content=mod_data.get("content", ""),
                    )
                    db.session.add(mod)
                    db.session.flush()

                    for q in mod_data.get("quiz", [])[:quiz_count]:
                        db.session.add(CourseQuiz(
                            module_id=mod.id,
                            question=q.get("question", ""),
                            options=json.dumps(q.get("options", [])),
                            correct_answer=int(q.get("correct", 0)),
                        ))

                    if include_assign:
                        asgn = mod_data.get("assignment") or {}
                        if asgn:
                            db.session.add(CourseAssignment(
                                module_id=mod.id,
                                title=asgn.get("title", "Assignment"),
                                description=asgn.get("description", ""),
                            ))

                    db.session.commit()
                    yield sse({"type": "module_done", "number": mod_num, "title": mod_title})

                yield sse({"type": "done", "id": course.id})

            except json.JSONDecodeError:
                db.session.rollback()
                yield sse({"type": "error", "msg": "AI returned invalid JSON. Please try again."})
            except Exception as e:
                db.session.rollback()
                yield sse({"type": "error", "msg": str(e)})

    return Response(
        stream_with_context(stream()),
        mimetype="text/event-stream",
        headers={"X-Accel-Buffering": "no", "Cache-Control": "no-cache"},
    )


@coursegen_bp.route("/<int:course_id>")
def view(course_id):
    course = Course.query.get_or_404(course_id)
    return render_template("coursegen/view.html", course=course)


@coursegen_bp.route("/<int:course_id>/delete", methods=["POST"])
def delete(course_id):
    course = Course.query.get_or_404(course_id)
    db.session.delete(course)
    db.session.commit()
    return jsonify({"success": True})


# ── Export: single module ─────────────────────────────────────────────────────

@coursegen_bp.route("/<int:course_id>/module/<int:module_id>/export/pdf")
def export_module_pdf(course_id, module_id):
    course = Course.query.get_or_404(course_id)
    mod    = CourseModule.query.filter_by(id=module_id, course_id=course_id).first_or_404()
    html   = build_module_html(course, mod)
    pdf    = html_to_pdf(html)
    fname  = f"Module_{mod.number}_{mod.title[:40].replace(' ','_')}.pdf"
    return Response(pdf, mimetype="application/pdf",
                    headers={"Content-Disposition": f'attachment; filename="{fname}"'})


@coursegen_bp.route("/<int:course_id>/module/<int:module_id>/export/docx")
def export_module_docx(course_id, module_id):
    course = Course.query.get_or_404(course_id)
    mod    = CourseModule.query.filter_by(id=module_id, course_id=course_id).first_or_404()
    buf    = build_module_docx(course, mod)
    fname  = f"Module_{mod.number}_{mod.title[:40].replace(' ','_')}.docx"
    return Response(
        buf.read(),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'},
    )


# ── Export: full course ───────────────────────────────────────────────────────

@coursegen_bp.route("/<int:course_id>/export/pdf")
def export_course_pdf(course_id):
    course = Course.query.get_or_404(course_id)
    html   = build_course_html(course)
    pdf    = html_to_pdf(html)
    fname  = f"{course.title[:60].replace(' ','_')}.pdf"
    return Response(pdf, mimetype="application/pdf",
                    headers={"Content-Disposition": f'attachment; filename="{fname}"'})


@coursegen_bp.route("/<int:course_id>/export/docx")
def export_course_docx(course_id):
    course = Course.query.get_or_404(course_id)
    buf    = build_course_docx(course)
    fname  = f"{course.title[:60].replace(' ','_')}.docx"
    return Response(
        buf.read(),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'},
    )
